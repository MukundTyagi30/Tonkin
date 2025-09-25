"""Document parser for SF84 Project Basis Reports."""

import os
import re
import logging
from pathlib import Path
from typing import Dict, List, Optional, Any
from dataclasses import dataclass, asdict
from datetime import datetime

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

from config.settings import SF84_HEADER_FIELDS, SF84_SECTIONS

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class SF84Document:
    """Structured representation of an SF84 Project Basis Report."""
    
    # File metadata
    file_path: str
    file_name: str
    file_size: int
    created_date: Optional[datetime] = None
    modified_date: Optional[datetime] = None
    
    # Header fields
    project_name: Optional[str] = None
    project_number: Optional[str] = None
    program_region: Optional[str] = None
    category: Optional[str] = None
    project_leader: Optional[str] = None
    project_reviewer: Optional[str] = None
    lead_disciplines: Optional[str] = None
    client: Optional[str] = None
    client_representative: Optional[str] = None
    
    # Document sections (text content)
    background: Optional[str] = None
    scope_of_work: Optional[str] = None
    scope_of_services: Optional[str] = None
    deliverables: Optional[str] = None
    reference_documents: Optional[str] = None
    existing_concept_design: Optional[str] = None
    assumptions: Optional[str] = None
    performance_requirements: Optional[str] = None
    operation_maintenance: Optional[str] = None
    monitoring_controls: Optional[str] = None
    
    # Trust scoring metadata
    trust_score: float = 0.0
    trust_badges: List[str] = None
    
    def __post_init__(self):
        if self.trust_badges is None:
            self.trust_badges = []
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for storage."""
        data = asdict(self)
        # Convert datetime objects to strings
        if self.created_date:
            data['created_date'] = self.created_date.isoformat()
        if self.modified_date:
            data['modified_date'] = self.modified_date.isoformat()
        return data
    
    def get_searchable_text(self) -> str:
        """Get all text content for semantic search."""
        text_parts = []
        
        # Add header information
        for field in ['project_name', 'project_number', 'program_region', 
                     'category', 'client', 'lead_disciplines']:
            value = getattr(self, field, None)
            if value:
                text_parts.append(value)
        
        # Add section content
        for field in ['background', 'scope_of_work', 'scope_of_services',
                     'deliverables', 'reference_documents', 'existing_concept_design',
                     'assumptions', 'performance_requirements', 'operation_maintenance',
                     'monitoring_controls']:
            value = getattr(self, field, None)
            if value:
                text_parts.append(value)
        
        return ' '.join(text_parts)


class DocumentParser:
    """Parser for SF84 Project Basis Reports."""
    
    def __init__(self):
        self.supported_extensions = ['.docx', '.pdf']
    
    def can_parse(self, file_path: str) -> bool:
        """Check if file can be parsed."""
        return Path(file_path).suffix.lower() in self.supported_extensions
    
    def parse_file(self, file_path: str) -> Optional[SF84Document]:
        """Parse a document file and extract structured data."""
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return None
        
        if not self.can_parse(file_path):
            logger.error(f"Unsupported file type: {file_path}")
            return None
        
        try:
            # Get file metadata
            file_stats = os.stat(file_path)
            file_path_obj = Path(file_path)
            
            doc = SF84Document(
                file_path=file_path,
                file_name=file_path_obj.name,
                file_size=file_stats.st_size,
                created_date=datetime.fromtimestamp(file_stats.st_ctime),
                modified_date=datetime.fromtimestamp(file_stats.st_mtime)
            )
            
            # Parse based on file type
            if file_path.lower().endswith('.docx'):
                self._parse_docx(file_path, doc)
            elif file_path.lower().endswith('.pdf'):
                self._parse_pdf(file_path, doc)
            
            # Calculate trust score
            self._calculate_trust_score(doc)
            
            return doc
            
        except Exception as e:
            logger.error(f"Error parsing {file_path}: {str(e)}")
            return None
    
    def _parse_docx(self, file_path: str, doc: SF84Document) -> None:
        """Parse DOCX file using python-docx."""
        if Document is None:
            raise ImportError("python-docx not available")
        
        docx_doc = Document(file_path)
        
        # Extract text from all paragraphs and tables
        all_text = []
        for paragraph in docx_doc.paragraphs:
            all_text.append(paragraph.text.strip())
        
        for table in docx_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text.append(cell.text.strip())
        
        full_text = '\n'.join(all_text)
        
        # Parse header fields and sections
        self._extract_header_fields(full_text, doc)
        self._extract_sections(full_text, doc)
    
    def _parse_pdf(self, file_path: str, doc: SF84Document) -> None:
        """Parse PDF file using PyPDF2."""
        if PyPDF2 is None:
            raise ImportError("PyPDF2 not available")
        
        text_content = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text_content.append(page.extract_text())
        
        full_text = '\n'.join(text_content)
        
        # Parse header fields and sections
        self._extract_header_fields(full_text, doc)
        self._extract_sections(full_text, doc)
    
    def _extract_header_fields(self, text: str, doc: SF84Document) -> None:
        """Extract header fields from document text."""
        # Mapping of field names to document attributes
        field_mapping = {
            "Project Name": "project_name",
            "Project Number": "project_number", 
            "Program/Region": "program_region",
            "Category": "category",
            "Project Leader": "project_leader",
            "Project Reviewer": "project_reviewer",
            "Lead Discipline(s)": "lead_disciplines",
            "Client": "client",
            "Client Representative": "client_representative"
        }
        
        for field_name, attr_name in field_mapping.items():
            value = self._extract_field_value(text, field_name)
            if value:
                setattr(doc, attr_name, value)
    
    def _extract_field_value(self, text: str, field_name: str) -> Optional[str]:
        """Extract value for a specific field from text."""
        # Try multiple patterns to find field values
        patterns = [
            rf"{re.escape(field_name)}\s*[:\-]\s*([^\n\r]+)",
            rf"{re.escape(field_name)}\s+([^\n\r]+)",
            rf"^{re.escape(field_name)}\s*([^\n\r]+)",
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.MULTILINE | re.IGNORECASE)
            if match:
                value = match.group(1).strip()
                # Clean up the value
                value = re.sub(r'\s+', ' ', value)  # Replace multiple spaces with single space
                if value and value != field_name:
                    return value
        
        return None
    
    def _extract_sections(self, text: str, doc: SF84Document) -> None:
        """Extract section content from document text."""
        section_mapping = {
            "Background": "background",
            "Scope of Work": "scope_of_work",
            "Scope of Services": "scope_of_services", 
            "Deliverables": "deliverables",
            "Reference documents & input data": "reference_documents",
            "Existing concept design": "existing_concept_design",
            "Assumptions": "assumptions",
            "Performance requirements": "performance_requirements",
            "Operation & maintenance": "operation_maintenance",
            "Monitoring & controls": "monitoring_controls"
        }
        
        for section_name, attr_name in section_mapping.items():
            content = self._extract_section_content(text, section_name)
            if content:
                setattr(doc, attr_name, content)
    
    def _extract_section_content(self, text: str, section_name: str) -> Optional[str]:
        """Extract content for a specific section."""
        # Try to find section by header
        section_pattern = rf"(?:^|\n)\s*{re.escape(section_name)}\s*[:\-]?\s*\n(.*?)(?=\n\s*(?:[A-Z][^:\n]*:|\d+\.|\w+\s+\w+\s*:|\Z))"
        
        match = re.search(section_pattern, text, re.DOTALL | re.IGNORECASE | re.MULTILINE)
        if match:
            content = match.group(1).strip()
            # Clean up the content
            content = re.sub(r'\n\s*\n', '\n', content)  # Remove extra blank lines
            content = re.sub(r'\s+', ' ', content)  # Normalize whitespace
            if content and len(content) > 10:  # Only return if meaningful content
                return content
        
        return None
    
    def _calculate_trust_score(self, doc: SF84Document) -> None:
        """Calculate trust score and badges for the document."""
        score = 0.0
        badges = []
        
        # Check for reviewer (25% weight)
        if doc.project_reviewer:
            score += 0.25
            badges.append("Has Reviewer")
        
        # Check header completeness (20% weight)
        header_fields = [doc.project_name, doc.project_number, doc.program_region,
                        doc.project_leader, doc.client]
        if all(field for field in header_fields):
            score += 0.20
            badges.append("Complete Header")
        
        # Check for recent document (15% weight)
        if doc.modified_date:
            days_old = (datetime.now() - doc.modified_date).days
            if days_old < 365:  # Less than 1 year old
                score += 0.15
                badges.append("Recent")
        
        # Check for content completeness (15% weight)
        content_sections = [doc.background, doc.scope_of_work, doc.deliverables]
        if all(section for section in content_sections):
            score += 0.15
            badges.append("Complete Content")
        
        # Check for region information (10% weight)
        if doc.program_region:
            score += 0.10
            badges.append("Region Specified")
        
        # Check for assumptions/requirements (10% weight)
        if doc.assumptions or doc.performance_requirements:
            score += 0.10
            badges.append("Includes Requirements")
        
        # Check for reference documents (5% weight)
        if doc.reference_documents:
            score += 0.05
            badges.append("References Cited")
        
        doc.trust_score = min(score, 1.0)  # Cap at 1.0
        doc.trust_badges = badges


def main():
    """Test the parser with sample files."""
    parser = DocumentParser()
    
    # Test with the sample file if it exists
    sample_file = "/Users/mukundtyagi/Desktop/tonkin protoype /SF84_Project_Basis_Report_V1.pdf"
    if os.path.exists(sample_file):
        print(f"Parsing: {sample_file}")
        doc = parser.parse_file(sample_file)
        if doc:
            print(f"Successfully parsed: {doc.file_name}")
            print(f"Trust Score: {doc.trust_score:.2f}")
            print(f"Trust Badges: {', '.join(doc.trust_badges)}")
            print(f"Searchable text length: {len(doc.get_searchable_text())} chars")
        else:
            print("Failed to parse document")
    else:
        print(f"Sample file not found: {sample_file}")


if __name__ == "__main__":
    main()
